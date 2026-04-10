"""Word document renderer for ensemble transcription results with Track Changes.

Renders ``EnsembleResult`` objects as a .docx file where disagreement points
appear as Track Changes.  The clinician opens the document in Word and uses
Accept/Reject Changes to resolve each disagreement.

Two rendering modes:

``"audit"`` (default):
    All disagreements shown as Track Changes for full traceability.
    MATCH        → plain run, no revision marks.
    VOXTRAL      → Whisper's word as tracked deletion, Voxtral's word as tracked insertion.
    WHISPER      → Voxtral's word as tracked deletion, Whisper's word as tracked insertion.
    CONTEXT_RULE → Voxtral's word as tracked deletion, context-derived word as tracked insertion.
    HUMAN_REVIEW → both words as tracked insertions, bold "[REVIEW]" marker prepended.

``"review"`` :
    Automated fixes baked in as normal text.  Only HUMAN_REVIEW items
    remain as Track Changes.  The clinician sees only the words that
    actually need their attention.
    MATCH        → plain run.
    VOXTRAL      → plain run (chosen word, automated fix accepted).
    WHISPER      → plain run (chosen word, automated fix accepted).
    CONTEXT_RULE → plain run (chosen word, automated fix accepted).
    HUMAN_REVIEW → both words as tracked insertions, bold "[REVIEW]" marker prepended.

Use ``render_ensemble_docx_pair`` to generate both documents in one call.
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx_revisions.paragraph import RevisionParagraph  # type: ignore[import-untyped]
from loguru import logger

from transcriber_radrx.ensemble.decision_rules import (
    DecisionSource,
    EnsembleResult,
    EnsembleWord,
)

# ---------------------------------------------------------------------------
# Custom exceptions
# ---------------------------------------------------------------------------


class DocxRendererError(RuntimeError):
    """Raised when the docx renderer encounters an unrecoverable error."""

    def __init__(self, message: str) -> None:
        super().__init__(message)


class OutputDirectoryMissingError(DocxRendererError):
    """Raised when the output directory for the .docx file does not exist."""

    def __init__(self, directory: Path) -> None:
        super().__init__(f"Output directory does not exist: {directory}")
        self.directory = directory


class DocxWriteError(DocxRendererError):
    """Raised when writing the .docx file to disk fails."""

    def __init__(self, path: Path, cause: Exception) -> None:
        super().__init__(f"Failed to write document to {path}: {cause}")
        self.path = path
        self.cause = cause


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------


def _add_review_marker(paragraph: RevisionParagraph) -> None:
    """Prepend a bold, red '[REVIEW]' marker run to *paragraph*.

    Args:
        paragraph: The paragraph to mark.
    """
    run = paragraph.add_run("[REVIEW] ")
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)


def _render_word(
    rp: RevisionParagraph,
    ew: EnsembleWord,
    author_voxtral: str,
    author_whisper: str,
    author_ensemble: str,
    *,
    mode: str = "audit",
    doc: Document | None = None,
) -> None:
    """Render one EnsembleWord into *rp* using appropriate Track Changes markup.

    The space after the word is always added as a plain run.

    Deletion approach: record ``rp.text`` length *before* adding the
    to-be-deleted word, then wrap it with ``add_tracked_deletion``.  This
    gives the correct character offsets relative to the current visible text
    in the paragraph, because ``RevisionParagraph.text`` excludes already-
    deleted runs while our newly added run is still visible.

    Args:
        rp: The paragraph (RevisionParagraph) to append to.
        ew: The ensemble word with provenance information.
        author_voxtral: Author name for Voxtral-sourced revisions.
        author_whisper: Author name for Whisper-sourced revisions.
        author_ensemble: Author name for ensemble-derived revisions.
        mode: ``"audit"`` renders all disagreements as Track Changes.
            ``"review"`` renders only HUMAN_REVIEW as Track Changes;
            automated fixes become plain text.
    """
    source = ew.source

    # In review mode, automated resolutions become plain text — UNLESS
    # the word is flagged for human review (e.g., context-inferred GyE
    # promotions where the promotion is uncertain).
    if (
        mode == "review"
        and not ew.needs_review
        and source
        in (
            DecisionSource.VOXTRAL,
            DecisionSource.WHISPER,
            DecisionSource.CONTEXT_RULE,
        )
    ):
        rp.add_run(ew.word)
        rp.add_run(" ")
        return

    if source == DecisionSource.MATCH:
        # Plain run — no revision marks.
        rp.add_run(ew.word)

    elif source == DecisionSource.VOXTRAL:
        # Whisper's word is deleted; Voxtral's word is inserted.
        whisper_word = ew.word_whisper or ""
        voxtral_word = ew.word or ""
        if whisper_word:
            offset_start = len(rp.text)
            rp.add_run(whisper_word)
            rp.add_tracked_deletion(offset_start, offset_start + len(whisper_word), author=author_whisper)
        rp.add_tracked_insertion(voxtral_word, author=author_ensemble)

    elif source == DecisionSource.WHISPER:
        # Voxtral's word is deleted; Whisper's word is inserted.
        voxtral_word = ew.word_voxtral or ""
        whisper_word = ew.word or ""
        if voxtral_word:
            offset_start = len(rp.text)
            rp.add_run(voxtral_word)
            rp.add_tracked_deletion(offset_start, offset_start + len(voxtral_word), author=author_voxtral)
        rp.add_tracked_insertion(whisper_word, author=author_ensemble)

    elif source == DecisionSource.CONTEXT_RULE:
        # Voxtral's word is deleted; the context-derived word is inserted.
        voxtral_word = ew.word_voxtral or ""
        context_word = ew.word or ""
        if voxtral_word:
            offset_start = len(rp.text)
            rp.add_run(voxtral_word)
            rp.add_tracked_deletion(offset_start, offset_start + len(voxtral_word), author=author_voxtral)
        rp.add_tracked_insertion(context_word, author=author_ensemble)

    elif source == DecisionSource.HUMAN_REVIEW:
        voxtral_word = ew.word_voxtral or ew.word or ""
        whisper_word = ew.word_whisper or ""

        if mode == "review" and doc is not None:
            # Review mode: render Voxtral's word as default text with a
            # comment listing both options.  The clinician either leaves
            # it (resolve comment = 1 click), edits it to the correct
            # word, or picks Whisper's suggestion from the comment.
            run = rp.add_run(voxtral_word)
            run.bold = True
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            options = [f"Voxtral: '{voxtral_word}'"]
            if whisper_word and whisper_word != voxtral_word:
                options.append(f"Whisper: '{whisper_word}'")
            comment_text = (
                "ASR disagreement — neither backend matched the vocabulary.\n"
                + "\n".join(options)
                + "\nEdit the highlighted word if needed, then resolve this comment."
            )
            doc.add_comment(runs=run, text=comment_text, author="Ensemble Review")
        else:
            # Audit mode: both words as tracked insertions with [REVIEW] marker.
            _add_review_marker(rp)
            if voxtral_word:
                rp.add_tracked_insertion(voxtral_word, author=author_voxtral)
            if whisper_word and whisper_word != voxtral_word:
                rp.add_run(" / ")
                rp.add_tracked_insertion(whisper_word, author=author_whisper)

    else:
        # Defensive fallback for any future DecisionSource values.
        logger.warning("Unknown DecisionSource {} — rendering as plain run", source)
        rp.add_run(ew.word)

    # Space after every word as a plain run.
    rp.add_run(" ")


def _render_summary_line(result: EnsembleResult) -> str:
    """Build a one-line summary string for an EnsembleResult.

    Args:
        result: The ensemble result to summarise.

    Returns:
        Summary string suitable for a document paragraph.
    """
    pct = round(result.agreement_rate * 100, 1)
    return (
        f"Agreement: {pct}% | "
        f"Voxtral chosen: {result.voxtral_chosen} | "
        f"Whisper chosen: {result.whisper_chosen} | "
        f"Context rules: {result.context_rule_count} | "
        f"Needs review: {result.review_count}"
    )


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def render_ensemble_docx(
    results: list[EnsembleResult],
    output_path: Path,
    *,
    title: str = "Ensemble Transcription Review",
    show_gold: bool = True,
    gold_texts: dict[tuple[str, str], str] | None = None,
    author_voxtral: str = "Voxtral Mini 3B",
    author_whisper: str = "Whisper large-v3",
    author_ensemble: str = "Ensemble",
    mode: str = "audit",
) -> Path:
    """Render ensemble results as a Word document with Track Changes.

    Each ``EnsembleResult`` becomes a section in the document.  Disagreement
    points (non-MATCH words) are rendered using Word Track Changes so the
    clinician can accept or reject each change in Word.

    Args:
        results: List of EnsembleResult objects (one per fixture × voice).
        output_path: Where to write the .docx file.
        title: Document title rendered on the cover page.
        show_gold: If True, include the gold reference text above each sample.
            Requires ``gold_texts`` to be provided; otherwise shows Voxtral raw.
        gold_texts: Optional mapping of (fixture_id, voice) → gold text.
            When provided and ``show_gold`` is True, the gold reference is shown
            as an italic line above the ensemble paragraph.
        author_voxtral: Author name used for Voxtral-sourced revision marks.
        author_whisper: Author name used for Whisper-sourced revision marks.
        author_ensemble: Author name used for ensemble/context-rule revision marks.
        mode: ``"audit"`` renders all disagreements as Track Changes for full
            traceability.  ``"review"`` bakes automated fixes into normal text
            and shows only HUMAN_REVIEW items as Track Changes — the clinician
            sees only what needs their attention.

    Returns:
        The ``output_path`` that was written.

    Raises:
        DocxRendererError: If the output directory does not exist or if
            writing fails.
    """
    output_path = Path(output_path)
    if not output_path.parent.exists():
        raise OutputDirectoryMissingError(output_path.parent)

    doc = Document()

    # ------------------------------------------------------------------
    # Cover / title section
    # ------------------------------------------------------------------
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(18)

    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_para.add_run(f"Generated: {date.today().isoformat()}  |  Samples: {len(results)}")

    total_words = sum(len(r.words) for r in results)
    needs_review_count = sum(1 for r in results if r.needs_review)
    total_review_words = sum(r.review_count for r in results)

    stats_para = doc.add_paragraph()
    stats_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    stats_para.add_run(
        f"Total words: {total_words}  |  "
        f"Samples needing review: {needs_review_count}  |  "
        f"Words flagged for review: {total_review_words}"
    )

    doc.add_paragraph()  # blank line

    # ------------------------------------------------------------------
    # Per-sample sections
    # ------------------------------------------------------------------
    for result in results:
        # Heading 2: fixture_id — voice
        heading_text = f"{result.fixture_id} — {result.voice}"
        doc.add_heading(heading_text, level=2)

        # Gold / reference text (italic)
        if show_gold:
            gold_para = doc.add_paragraph()
            gold_run = gold_para.add_run("")
            gold_run.italic = True
            gold_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            key = (result.fixture_id, result.voice)
            if gold_texts and key in gold_texts:
                gold_run.text = f"Gold: {gold_texts[key]}"
            else:
                # Fall back to Voxtral raw as approximate reference
                gold_run.text = f"Voxtral raw: {result.text_voxtral}"

        # Ensemble paragraph with Track Changes
        ensemble_para = doc.add_paragraph()
        rp = RevisionParagraph.from_paragraph(ensemble_para)

        for ew in result.words:
            try:
                _render_word(
                    rp,
                    ew,
                    author_voxtral,
                    author_whisper,
                    author_ensemble,
                    mode=mode,
                    doc=doc,
                )
            except Exception:
                logger.exception(
                    "Failed to render word fixture={} voice={} word={}",
                    result.fixture_id,
                    result.voice,
                    ew.word,
                )
                # Fallback: plain run
                rp.add_run(ew.word + " ")

        # Summary line
        summary_para = doc.add_paragraph()
        summary_run = summary_para.add_run(_render_summary_line(result))
        summary_run.italic = True
        summary_run.font.size = Pt(9)
        summary_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_paragraph()  # blank separator between samples

    # ------------------------------------------------------------------
    # Write the document
    # ------------------------------------------------------------------
    try:
        doc.save(str(output_path))
    except Exception as exc:
        raise DocxWriteError(output_path, exc) from exc

    logger.info("Ensemble docx ({} mode) written to {}", mode, output_path)
    return output_path


def render_ensemble_docx_pair(
    results: list[EnsembleResult],
    audit_path: Path,
    review_path: Path,
    **kwargs: object,
) -> tuple[Path, Path]:
    """Generate both audit and review documents from the same ensemble results.

    The **audit** document shows every ensemble decision as a Track Change
    for full traceability.  The **review** document bakes automated fixes
    into normal text and shows only the words flagged for human review as
    Track Changes.

    Args:
        results: List of EnsembleResult objects.
        audit_path: Output path for the audit document.
        review_path: Output path for the review document.
        **kwargs: Passed through to ``render_ensemble_docx`` (title,
            show_gold, gold_texts, author_voxtral, author_whisper,
            author_ensemble).

    Returns:
        Tuple of (audit_path, review_path).
    """
    audit_title = str(kwargs.pop("title", "Ensemble Transcription Review"))
    render_ensemble_docx(
        results,
        audit_path,
        mode="audit",
        title=f"{audit_title} — Audit",
        **kwargs,  # type: ignore[arg-type]
    )
    render_ensemble_docx(
        results,
        review_path,
        mode="review",
        title=f"{audit_title} — Review",
        **kwargs,  # type: ignore[arg-type]
    )
    return audit_path, review_path
