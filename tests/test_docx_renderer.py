"""Tests for the ensemble docx renderer with Track Changes.

Tests verify:
- A valid .docx file is produced (can be re-read by python-docx).
- MATCH words produce plain runs with no revision marks.
- VOXTRAL-source words produce w:del (Whisper's word) + w:ins (Voxtral's word).
- WHISPER-source words produce w:del (Voxtral's word) + w:ins (Whisper's word).
- CONTEXT_RULE words produce w:del + w:ins (author=ensemble).
- HUMAN_REVIEW words produce w:ins for both words and a "[REVIEW]" marker.
- The document contains the correct number of headings.
- show_gold=True and show_gold=False both work.
- gold_texts mapping is honoured when provided.
- Integration test on real EnsembleResults built from particle corpus data.
- DocxRendererError is raised when output directory does not exist.
"""

from __future__ import annotations

import json
import re
from pathlib import Path

import pytest
from docx import Document  # type: ignore[import-untyped]
from docx.document import Document as DocxDocument  # type: ignore[import-untyped]

from transcriber_radrx.ensemble.decision_rules import (
    DecisionSource,
    EnsembleResult,
    EnsembleWord,
    ensemble_transcriptions,
)
from transcriber_radrx.ensemble.docx_renderer import DocxRendererError, render_ensemble_docx

# ---------------------------------------------------------------------------
# Fixtures — synthetic EnsembleWords and EnsembleResults
# ---------------------------------------------------------------------------

_REPO_ROOT = Path(__file__).parents[1]
_REPORTS_DIR = _REPO_ROOT / "tests" / "validation" / "reports"
_VOCAB_FILE = _REPO_ROOT / "data" / "rt_vocabulary.txt"


def _make_match_word(word: str) -> EnsembleWord:
    """Create a MATCH EnsembleWord."""
    return EnsembleWord(
        word=word,
        source=DecisionSource.MATCH,
        word_voxtral=word,
        word_whisper=word,
        rule_id=None,
        confidence=1.0,
        needs_review=False,
    )


def _make_voxtral_word(voxtral: str, whisper: str) -> EnsembleWord:
    """Create a VOXTRAL-source EnsembleWord."""
    return EnsembleWord(
        word=voxtral,
        source=DecisionSource.VOXTRAL,
        word_voxtral=voxtral,
        word_whisper=whisper,
        rule_id="FORMATTING_DEFAULT",
        confidence=0.7,
        needs_review=False,
    )


def _make_whisper_word(whisper: str, voxtral: str) -> EnsembleWord:
    """Create a WHISPER-source EnsembleWord."""
    return EnsembleWord(
        word=whisper,
        source=DecisionSource.WHISPER,
        word_voxtral=voxtral,
        word_whisper=whisper,
        rule_id="VOCABULARY_MATCH",
        confidence=0.9,
        needs_review=False,
    )


def _make_context_word(chosen: str, voxtral: str, whisper: str) -> EnsembleWord:
    """Create a CONTEXT_RULE EnsembleWord."""
    return EnsembleWord(
        word=chosen,
        source=DecisionSource.CONTEXT_RULE,
        word_voxtral=voxtral,
        word_whisper=whisper,
        rule_id="DOSE_UNIT_CONTEXT",
        confidence=0.85,
        needs_review=True,
    )


def _make_review_word(voxtral: str, whisper: str) -> EnsembleWord:
    """Create a HUMAN_REVIEW EnsembleWord."""
    return EnsembleWord(
        word=voxtral,
        source=DecisionSource.HUMAN_REVIEW,
        word_voxtral=voxtral,
        word_whisper=whisper,
        rule_id="BOTH_WRONG",
        confidence=0.3,
        needs_review=True,
    )


def _make_result(
    fixture_id: str,
    voice: str,
    words: list[EnsembleWord],
) -> EnsembleResult:
    """Build an EnsembleResult from a list of EnsembleWords."""
    text_ensemble = " ".join(w.word for w in words if w.word)
    text_voxtral = " ".join(w.word_voxtral or w.word for w in words if w.word)
    text_whisper = " ".join(w.word_whisper or w.word for w in words if w.word)
    review_count = sum(1 for w in words if w.needs_review)
    voxtral_chosen = sum(1 for w in words if w.source in (DecisionSource.VOXTRAL, DecisionSource.MATCH))
    whisper_chosen = sum(1 for w in words if w.source == DecisionSource.WHISPER)
    context_rule_count = sum(1 for w in words if w.source == DecisionSource.CONTEXT_RULE)
    return EnsembleResult(
        fixture_id=fixture_id,
        voice=voice,
        text_voxtral=text_voxtral,
        text_whisper=text_whisper,
        text_ensemble=text_ensemble,
        words=words,
        needs_review=any(w.needs_review for w in words),
        review_count=review_count,
        agreement_rate=voxtral_chosen / len(words) if words else 1.0,
        voxtral_chosen=voxtral_chosen,
        whisper_chosen=whisper_chosen,
        context_rule_count=context_rule_count,
    )


@pytest.fixture()
def simple_result() -> EnsembleResult:
    """EnsembleResult with one word of each source type."""
    words = [
        _make_match_word("Prescribed"),
        _make_voxtral_word("55.8", "fifty"),
        _make_whisper_word("GyE", "Gy"),
        _make_context_word("GyE", "Gy", "Jai"),
        _make_review_word("craniofacial", "craniophore"),
        _make_match_word("fractions."),
    ]
    return _make_result("test-0001", "en_US-lessac-high", words)


@pytest.fixture()
def all_match_result() -> EnsembleResult:
    """EnsembleResult where every word is a MATCH."""
    words = [_make_match_word(w) for w in ["the", "quick", "brown", "fox"]]
    return _make_result("test-match", "en_US-lessac-high", words)


# ---------------------------------------------------------------------------
# Helpers for reading back the output
# ---------------------------------------------------------------------------


def _para_xml(doc: DocxDocument, paragraph_index: int) -> str:
    """Return the XML of a paragraph by index."""
    para = doc.paragraphs[paragraph_index]
    xml: str = para._element.xml
    return xml


def _all_xml(doc: DocxDocument) -> str:
    """Concatenate all paragraph XML in the document."""
    parts: list[str] = []
    for p in doc.paragraphs:
        xml: str = p._element.xml
        parts.append(xml)
    return "\n".join(parts)


def _count_headings(doc: DocxDocument, level: int) -> int:
    """Count paragraphs with the given heading style level."""
    style_name = f"Heading {level}"
    count = 0
    for p in doc.paragraphs:
        style = p.style
        if style is not None and style.name == style_name:
            count += 1
    return count


# ---------------------------------------------------------------------------
# Basic contract tests
# ---------------------------------------------------------------------------


class TestRenderEnsembleDocxBasic:
    """Basic rendering contracts."""

    def test_produces_valid_docx(self, tmp_path: Path, simple_result: EnsembleResult) -> None:
        """Contract: render_ensemble_docx produces a file readable by python-docx."""
        out = tmp_path / "test.docx"
        result_path = render_ensemble_docx([simple_result], out)
        assert result_path == out
        assert out.exists()
        # Verify it can be opened by python-docx
        doc = Document(str(out))
        assert len(doc.paragraphs) > 0

    def test_returns_output_path(self, tmp_path: Path, simple_result: EnsembleResult) -> None:
        """Contract: render_ensemble_docx returns the output_path."""
        out = tmp_path / "check_return.docx"
        returned = render_ensemble_docx([simple_result], out)
        assert returned == out

    def test_raises_on_missing_directory(self, simple_result: EnsembleResult) -> None:
        """Contract: DocxRendererError raised when output directory does not exist."""
        bad_path = Path("/nonexistent_dir_xyz/output.docx")
        with pytest.raises(DocxRendererError):
            render_ensemble_docx([simple_result], bad_path)

    def test_empty_results_list(self, tmp_path: Path) -> None:
        """Contract: empty results list produces a document with just the title."""
        out = tmp_path / "empty.docx"
        render_ensemble_docx([], out)
        doc = Document(str(out))
        assert len(doc.paragraphs) > 0  # title paragraph at minimum

    def test_one_heading_per_sample(self, tmp_path: Path, simple_result: EnsembleResult) -> None:
        """Contract: one Heading 2 per EnsembleResult."""
        out = tmp_path / "headings.docx"
        results = [
            simple_result,
            _make_result("test-0002", "en_GB-alan-medium", [_make_match_word("hello")]),
            _make_result("test-0003", "en_US-lessac-high", [_make_match_word("world")]),
        ]
        render_ensemble_docx(results, out)
        doc = Document(str(out))
        assert _count_headings(doc, 2) == 3

    def test_heading_text_contains_fixture_and_voice(self, tmp_path: Path, simple_result: EnsembleResult) -> None:
        """Contract: heading text is '{fixture_id} — {voice}'."""
        out = tmp_path / "heading_text.docx"
        render_ensemble_docx([simple_result], out)
        doc = Document(str(out))
        headings = [p for p in doc.paragraphs if p.style.name == "Heading 2"]
        assert len(headings) == 1
        assert simple_result.fixture_id in headings[0].text
        assert simple_result.voice in headings[0].text


# ---------------------------------------------------------------------------
# Track Changes XML structure tests
# ---------------------------------------------------------------------------


class TestTrackChangesXmlStructure:
    """Tests that verify the actual XML contains the correct revision marks."""

    def test_match_words_have_no_revision_marks(self, tmp_path: Path, all_match_result: EnsembleResult) -> None:
        """Contract: MATCH words produce no w:ins or w:del elements."""
        out = tmp_path / "match.docx"
        render_ensemble_docx([all_match_result], out)
        doc = Document(str(out))
        all_xml = _all_xml(doc)
        assert "<w:ins" not in all_xml
        assert "<w:del" not in all_xml

    def test_voxtral_source_has_del_and_ins(self, tmp_path: Path) -> None:
        """Contract: VOXTRAL-source word produces w:del (Whisper's) + w:ins (Voxtral's)."""
        words = [_make_voxtral_word("55.8", "fifty")]
        result = _make_result("vox-test", "voice", words)
        out = tmp_path / "voxtral_word.docx"
        render_ensemble_docx([result], out)
        doc = Document(str(out))
        all_xml = _all_xml(doc)
        assert "<w:del" in all_xml
        assert "<w:ins" in all_xml
        # The deleted text should be Whisper's word
        assert "fifty" in all_xml
        # The inserted text should be Voxtral's word
        assert "55.8" in all_xml

    def test_voxtral_deletion_author_is_whisper(self, tmp_path: Path) -> None:
        """Contract: VOXTRAL-source deletion is attributed to the Whisper author."""
        words = [_make_voxtral_word("55.8", "fifty")]
        result = _make_result("vox-author", "voice", words)
        out = tmp_path / "voxtral_author.docx"
        render_ensemble_docx([result], out, author_whisper="WhisperTestAuthor")
        doc = Document(str(out))
        all_xml = _all_xml(doc)
        # The w:del element should name WhisperTestAuthor
        del_match = re.search(r'<w:del[^>]*w:author="([^"]+)"', all_xml)
        assert del_match is not None
        assert del_match.group(1) == "WhisperTestAuthor"

    def test_whisper_source_has_del_and_ins(self, tmp_path: Path) -> None:
        """Contract: WHISPER-source word produces w:del (Voxtral's) + w:ins (Whisper's)."""
        words = [_make_whisper_word("GyE", "Gy")]
        result = _make_result("whi-test", "voice", words)
        out = tmp_path / "whisper_word.docx"
        render_ensemble_docx([result], out)
        doc = Document(str(out))
        all_xml = _all_xml(doc)
        assert "<w:del" in all_xml
        assert "<w:ins" in all_xml
        assert "Gy" in all_xml  # Voxtral's word (deleted)
        assert "GyE" in all_xml  # Whisper's word (inserted)

    def test_whisper_deletion_author_is_voxtral(self, tmp_path: Path) -> None:
        """Contract: WHISPER-source deletion is attributed to the Voxtral author."""
        words = [_make_whisper_word("GyE", "Gy")]
        result = _make_result("whi-author", "voice", words)
        out = tmp_path / "whisper_author.docx"
        render_ensemble_docx([result], out, author_voxtral="VoxtralTestAuthor")
        doc = Document(str(out))
        all_xml = _all_xml(doc)
        del_match = re.search(r'<w:del[^>]*w:author="([^"]+)"', all_xml)
        assert del_match is not None
        assert del_match.group(1) == "VoxtralTestAuthor"

    def test_context_rule_has_del_and_ins(self, tmp_path: Path) -> None:
        """Contract: CONTEXT_RULE word produces w:del (Voxtral's) + w:ins (context word)."""
        words = [_make_context_word("GyE", "Gy", "Jai")]
        result = _make_result("ctx-test", "voice", words)
        out = tmp_path / "context_word.docx"
        render_ensemble_docx([result], out)
        doc = Document(str(out))
        all_xml = _all_xml(doc)
        assert "<w:del" in all_xml
        assert "<w:ins" in all_xml
        # Voxtral's word should be deleted
        assert "Gy" in all_xml

    def test_context_rule_insertion_author_is_ensemble(self, tmp_path: Path) -> None:
        """Contract: CONTEXT_RULE insertion is attributed to the ensemble author."""
        words = [_make_context_word("GyE", "Gy", "Jai")]
        result = _make_result("ctx-author", "voice", words)
        out = tmp_path / "context_author.docx"
        render_ensemble_docx([result], out, author_ensemble="EnsembleTestAuthor")
        doc = Document(str(out))
        all_xml = _all_xml(doc)
        ins_match = re.search(r'<w:ins[^>]*w:author="([^"]+)"', all_xml)
        assert ins_match is not None
        assert ins_match.group(1) == "EnsembleTestAuthor"

    def test_human_review_has_review_marker(self, tmp_path: Path) -> None:
        """Contract: HUMAN_REVIEW words have a '[REVIEW]' text marker."""
        words = [_make_review_word("craniofacial", "craniophore")]
        result = _make_result("rev-test", "voice", words)
        out = tmp_path / "review_word.docx"
        render_ensemble_docx([result], out)
        doc = Document(str(out))
        all_xml = _all_xml(doc)
        assert "[REVIEW]" in all_xml

    def test_human_review_has_insertions_for_both_words(self, tmp_path: Path) -> None:
        """Contract: HUMAN_REVIEW words show both backend words as tracked insertions."""
        words = [_make_review_word("craniofacial", "craniophore")]
        result = _make_result("rev-ins", "voice", words)
        out = tmp_path / "review_ins.docx"
        render_ensemble_docx([result], out)
        doc = Document(str(out))
        all_xml = _all_xml(doc)
        # Both words should appear (at least one as a tracked insertion)
        assert "craniofacial" in all_xml
        assert "craniophore" in all_xml
        # At least one w:ins
        assert "<w:ins" in all_xml

    def test_mixed_sources_have_both_del_and_ins(self, tmp_path: Path, simple_result: EnsembleResult) -> None:
        """Contract: a result with mixed sources has both w:del and w:ins."""
        out = tmp_path / "mixed.docx"
        render_ensemble_docx([simple_result], out)
        doc = Document(str(out))
        all_xml = _all_xml(doc)
        assert "<w:del" in all_xml
        assert "<w:ins" in all_xml


# ---------------------------------------------------------------------------
# show_gold tests
# ---------------------------------------------------------------------------


class TestShowGold:
    """Tests for the show_gold and gold_texts options."""

    def test_show_gold_true_uses_gold_texts_when_provided(self, tmp_path: Path, simple_result: EnsembleResult) -> None:
        """Contract: gold_texts mapping is rendered as 'Gold: ...' when provided."""
        gold = {(simple_result.fixture_id, simple_result.voice): "Gold reference text here."}
        out = tmp_path / "gold_true.docx"
        render_ensemble_docx([simple_result], out, show_gold=True, gold_texts=gold)
        doc = Document(str(out))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Gold reference text here" in all_text

    def test_show_gold_true_falls_back_to_voxtral_raw_when_no_gold_texts(
        self, tmp_path: Path, simple_result: EnsembleResult
    ) -> None:
        """Contract: without gold_texts, show_gold=True shows Voxtral raw text."""
        out = tmp_path / "gold_fallback.docx"
        render_ensemble_docx([simple_result], out, show_gold=True)
        doc = Document(str(out))
        all_text = " ".join(p.text for p in doc.paragraphs)
        # Should contain the Voxtral raw text label
        assert "Voxtral raw" in all_text

    def test_show_gold_false_omits_gold_paragraph(self, tmp_path: Path, simple_result: EnsembleResult) -> None:
        """Contract: show_gold=False omits the gold/voxtral reference paragraph."""
        gold = {(simple_result.fixture_id, simple_result.voice): "Gold reference text here."}
        out = tmp_path / "gold_false.docx"
        render_ensemble_docx([simple_result], out, show_gold=False, gold_texts=gold)
        doc = Document(str(out))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Gold reference text here" not in all_text
        assert "Voxtral raw" not in all_text

    def test_show_gold_default_is_true(self, tmp_path: Path, simple_result: EnsembleResult) -> None:
        """Contract: show_gold defaults to True."""
        out = tmp_path / "gold_default.docx"
        render_ensemble_docx([simple_result], out)
        doc = Document(str(out))
        all_text = " ".join(p.text for p in doc.paragraphs)
        # Default is True, so there should be a gold/voxtral line
        assert "Voxtral raw" in all_text or "Gold" in all_text


# ---------------------------------------------------------------------------
# Summary line tests
# ---------------------------------------------------------------------------


class TestSummaryLine:
    """Tests for the per-sample summary line."""

    def test_summary_line_contains_agreement_percentage(self, tmp_path: Path, simple_result: EnsembleResult) -> None:
        """Contract: the summary line contains 'Agreement:'."""
        out = tmp_path / "summary.docx"
        render_ensemble_docx([simple_result], out)
        doc = Document(str(out))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Agreement:" in all_text

    def test_summary_line_contains_needs_review_count(self, tmp_path: Path, simple_result: EnsembleResult) -> None:
        """Contract: the summary line contains 'Needs review:'."""
        out = tmp_path / "summary_review.docx"
        render_ensemble_docx([simple_result], out)
        doc = Document(str(out))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Needs review:" in all_text


# ---------------------------------------------------------------------------
# Integration test with real particle corpus data
# ---------------------------------------------------------------------------


class TestIntegrationWithParticleCorpus:
    """Integration tests using real EnsembleResults from the particle corpus."""

    @pytest.fixture(autouse=True)
    def _check_data_files(self) -> None:
        """Skip if required bake-off JSON files are not present."""
        voxtral_json = _REPORTS_DIR / "bakeoff_particle_voxtral_2026-04-09.json"
        whisper_json = _REPORTS_DIR / "bakeoff_particle_whisper_medasr_2026-04-09.json"
        vocab_file = _VOCAB_FILE
        if not (voxtral_json.exists() and whisper_json.exists() and vocab_file.exists()):
            pytest.skip("Particle corpus bake-off data not available")

    def _load_vocab(self) -> set[str]:
        """Load RT vocabulary terms."""
        terms: set[str] = set()
        with _VOCAB_FILE.open(encoding="utf-8") as fh:
            for line in fh:
                stripped = line.strip()
                if stripped and not stripped.startswith("#"):
                    terms.add(stripped.lower())
        return terms

    def _extract_samples(self, report: dict, backend_name: str) -> dict[str, dict[str, dict]]:
        """Extract samples indexed by voice → fixture_id."""
        index: dict[str, dict[str, dict]] = {}
        results = report.get("results", [])
        for be in results:
            if not isinstance(be, dict) or be.get("backend") != backend_name:
                continue
            for ve in be.get("by_voice", []):
                if not isinstance(ve, dict):
                    continue
                voice = str(ve.get("voice", ""))
                vm: dict[str, dict] = {}
                for s in ve.get("samples", []):
                    if isinstance(s, dict):
                        vm[str(s.get("fixture_id", ""))] = s
                index[voice] = vm
        return index

    def test_render_at_least_three_real_results(self, tmp_path: Path) -> None:
        """Contract: renders at least 3 real EnsembleResults from particle corpus without errors."""
        voxtral_path = _REPORTS_DIR / "bakeoff_particle_voxtral_2026-04-09.json"
        whisper_path = _REPORTS_DIR / "bakeoff_particle_whisper_medasr_2026-04-09.json"

        with voxtral_path.open() as f:
            vox_report = json.load(f)
        with whisper_path.open() as f:
            whi_report = json.load(f)

        vocab = self._load_vocab()
        vox_by_voice = self._extract_samples(vox_report, "voxtral")
        whi_by_voice = self._extract_samples(whi_report, "mlx_whisper")
        common_voices = sorted(set(vox_by_voice.keys()) & set(whi_by_voice.keys()))

        results: list[EnsembleResult] = []
        gold_texts: dict[tuple[str, str], str] = {}

        for voice in common_voices[:1]:  # Only first voice to keep test fast
            vox_fixtures = vox_by_voice[voice]
            whi_fixtures = whi_by_voice[voice]
            common_fixtures = sorted(set(vox_fixtures.keys()) & set(whi_fixtures.keys()))

            for fid in common_fixtures[:3]:  # First 3 fixtures
                vox_s = vox_fixtures[fid]
                whi_s = whi_fixtures[fid]
                text_voxtral = str(vox_s.get("raw_transcription", ""))
                text_whisper = str(whi_s.get("raw_transcription", ""))
                ground_truth = str(vox_s.get("ground_truth", ""))

                er = ensemble_transcriptions(
                    text_voxtral=text_voxtral,
                    text_whisper=text_whisper,
                    vocabulary=vocab,
                    fixture_id=fid,
                    voice=voice,
                )
                results.append(er)
                if ground_truth:
                    gold_texts[(fid, voice)] = ground_truth

        assert len(results) >= 3, f"Expected at least 3 results, got {len(results)}"

        out = tmp_path / "integration_test.docx"
        render_ensemble_docx(results, out, gold_texts=gold_texts)

        assert out.exists()
        doc = Document(str(out))
        assert _count_headings(doc, 2) == len(results)

    def test_real_results_have_correct_xml_structure(self, tmp_path: Path) -> None:
        """Contract: real EnsembleResults produce valid Track Changes XML when there are disagreements."""
        voxtral_path = _REPORTS_DIR / "bakeoff_particle_voxtral_2026-04-09.json"
        whisper_path = _REPORTS_DIR / "bakeoff_particle_whisper_medasr_2026-04-09.json"

        with voxtral_path.open() as f:
            vox_report = json.load(f)
        with whisper_path.open() as f:
            whi_report = json.load(f)

        vocab = self._load_vocab()
        vox_by_voice = self._extract_samples(vox_report, "voxtral")
        whi_by_voice = self._extract_samples(whi_report, "mlx_whisper")
        common_voices = sorted(set(vox_by_voice.keys()) & set(whi_by_voice.keys()))

        # Find a result with at least one disagreement
        disagreement_result: EnsembleResult | None = None
        for voice in common_voices:
            vox_fixtures = vox_by_voice[voice]
            whi_fixtures = whi_by_voice[voice]
            for fid in sorted(set(vox_fixtures.keys()) & set(whi_fixtures.keys())):
                er = ensemble_transcriptions(
                    text_voxtral=str(vox_fixtures[fid].get("raw_transcription", "")),
                    text_whisper=str(whi_fixtures[fid].get("raw_transcription", "")),
                    vocabulary=vocab,
                    fixture_id=fid,
                    voice=voice,
                )
                if er.voxtral_chosen != len(er.words) and er.whisper_chosen > 0:
                    disagreement_result = er
                    break
            if disagreement_result:
                break

        if disagreement_result is None:
            pytest.skip("No disagreement found in particle corpus (all MATCH)")

        out = tmp_path / "xml_check.docx"
        render_ensemble_docx([disagreement_result], out)
        doc = Document(str(out))
        all_xml = _all_xml(doc)

        # Should have revision marks since there are disagreements
        assert "<w:ins" in all_xml or "<w:del" in all_xml, "Expected revision marks for disagreement result"
